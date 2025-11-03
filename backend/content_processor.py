"""
Advanced Content Processing Module for Custom Topic Upload
Supports multiple file formats and content extraction methods
"""

import os
import re
import json
import requests
from typing import Dict, List, Any, Optional
from datetime import datetime
import PyPDF2
import pdfplumber
import docx
from bs4 import BeautifulSoup
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    print("Warning: python-magic not available. File type detection will be limited.")
import hashlib
import google.generativeai as genai
import logging

logger = logging.getLogger(__name__)

class ContentProcessor:
    """Advanced content processor for multiple file formats and sources"""
    
    def __init__(self, upload_dir: str = "uploads"):
        """Initialize content processor with upload directory"""
        self.upload_dir = upload_dir
        self.ensure_upload_directory()
        
        # Supported file types
        self.supported_formats = {
            'text': ['.txt', '.md', '.rst'],
            'document': ['.pdf', '.docx', '.doc'],
            'data': ['.json', '.csv', '.xml'],
            'web': ['http://', 'https://']
        }
        
        # Maximum file size (10MB)
        self.max_file_size = 10 * 1024 * 1024
        
        # Configure Google Gemini AI
        self.gemini_api_key = os.getenv('GEMINI_API_KEY')
        if self.gemini_api_key:
            genai.configure(api_key=self.gemini_api_key)
            self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
            logger.info("✅ Google Gemini AI configured for PDF question generation")
        else:
            self.gemini_model = None
            logger.warning("⚠️ GEMINI_API_KEY not found. PDF question generation will be limited.")
        
        print("📁 Content Processor initialized with advanced file support")
    
    def ensure_upload_directory(self):
        """Ensure upload directory exists"""
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir)
            print(f"📂 Created upload directory: {self.upload_dir}")
    
    def validate_file(self, file_path: str, file_size: int) -> Dict[str, Any]:
        """Validate uploaded file"""
        validation_result = {
            'is_valid': False,
            'file_type': None,
            'error_message': None,
            'file_info': {}
        }
        
        # Check file size
        if file_size > self.max_file_size:
            validation_result['error_message'] = f"File too large. Maximum size: {self.max_file_size // (1024*1024)}MB"
            return validation_result
        
        # Check file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        file_type = None
        
        for format_type, extensions in self.supported_formats.items():
            if format_type != 'web' and file_ext in extensions:
                file_type = format_type
                break
        
        if not file_type:
            validation_result['error_message'] = f"Unsupported file format: {file_ext}"
            return validation_result
        
        # Detect MIME type for security
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_file(file_path, mime=True)
                validation_result['file_info']['mime_type'] = mime_type
            except:
                pass  # magic library might not be available
        
        validation_result['is_valid'] = True
        validation_result['file_type'] = file_type
        validation_result['file_info']['extension'] = file_ext
        validation_result['file_info']['size'] = file_size
        
        return validation_result
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Clean and normalize text
            content = re.sub(r'\n\s*\n', '\n\n', content)  # Normalize line breaks
            content = re.sub(r'\s+', ' ', content)  # Normalize spaces
            
            return content.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
    
    def extract_text_from_pdf(self, file_path: str, use_advanced: bool = True) -> str:
        """
        Extract text from PDF files with advanced processing
        
        Args:
            file_path: Path to PDF file
            use_advanced: Use pdfplumber for better text extraction
        
        Returns:
            Extracted text content
        """
        try:
            if use_advanced:
                # Use pdfplumber for better text extraction with layout preservation
                text_content = []
                
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text with layout
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            text_content.append(f"=== Page {page_num + 1} ===\n{page_text}")
                        
                        # Extract tables if present
                        tables = page.extract_tables()
                        if tables:
                            for table_num, table in enumerate(tables):
                                table_text = f"\n--- Table {table_num + 1} on Page {page_num + 1} ---\n"
                                for row in table:
                                    if row:
                                        # Filter out None values and join cells
                                        row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                        table_text += row_text + "\n"
                                text_content.append(table_text)
                
                content = '\n\n'.join(text_content)
            else:
                # Fallback to PyPDF2
                text_content = []
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"Page {page_num + 1}:\n{page_text}")
                
                content = '\n\n'.join(text_content)
            
            # Clean extracted text
            content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
            content = re.sub(r'[ \t]+', ' ', content)  # Normalize spaces but keep line breaks
            
            logger.info(f"✅ Extracted {len(content)} characters from PDF")
            return content.strip()
        
        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {str(e)}")
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            content = '\n\n'.join(text_content)
            
            # Clean text
            content = re.sub(r'\n\s*\n', '\n\n', content)
            content = re.sub(r'\s+', ' ', content)
            
            return content.strip()
        
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def extract_text_from_json(self, file_path: str) -> str:
        """Extract and format text from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            def json_to_text(obj, level=0):
                indent = "  " * level
                text_parts = []
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text_parts.append(f"{indent}{key}:")
                            text_parts.append(json_to_text(value, level + 1))
                        else:
                            text_parts.append(f"{indent}{key}: {value}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        text_parts.append(f"{indent}Item {i + 1}:")
                        text_parts.append(json_to_text(item, level + 1))
                else:
                    text_parts.append(f"{indent}{obj}")
                
                return '\n'.join(text_parts)
            
            return json_to_text(data)
        
        except Exception as e:
            raise Exception(f"JSON extraction failed: {str(e)}")
    
    def extract_text_from_csv(self, file_path: str) -> str:
        """Extract and format text from CSV files"""
        try:
            import csv
            text_content = []
            
            with open(file_path, 'r', encoding='utf-8') as file:
                # Try to detect delimiter
                sample = file.read(1024)
                file.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.reader(file, delimiter=delimiter)
                headers = next(reader, None)
                
                if headers:
                    text_content.append("Headers: " + " | ".join(headers))
                    text_content.append("")
                
                # Read rows (limit to first 100 for processing efficiency)
                for i, row in enumerate(reader):
                    if i >= 100:  # Limit rows
                        text_content.append(f"... (showing first 100 rows of {i+1}+ total)")
                        break
                    if row:
                        text_content.append(" | ".join(str(cell) for cell in row))
            
            return '\n'.join(text_content)
        
        except Exception as e:
            raise Exception(f"CSV extraction failed: {str(e)}")
    
    def extract_content_from_url(self, url: str) -> str:
        """Extract text content from web URLs"""
        try:
            # Validate URL
            if not (url.startswith('http://') or url.startswith('https://')):
                raise Exception("Invalid URL format")
            
            # Fetch content with timeout
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            # Clean text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Limit content length for processing
            if len(text) > 10000:
                text = text[:10000] + "...\n[Content truncated for processing]"
            
            return text
        
        except Exception as e:
            raise Exception(f"URL extraction failed: {str(e)}")
    
    def process_content(self, source: str, source_type: str = "auto") -> Dict[str, Any]:
        """
        Main method to process content from various sources
        
        Args:
            source: File path, URL, or direct text content
            source_type: "file", "url", "text", or "auto" for auto-detection
        
        Returns:
            Dictionary with extracted content and metadata
        """
        
        processing_result = {
            'success': False,
            'content': '',
            'metadata': {
                'source_type': source_type,
                'processed_at': datetime.now().isoformat(),
                'content_hash': None,
                'word_count': 0,
                'character_count': 0
            },
            'error': None
        }
        
        try:
            # Auto-detect source type
            if source_type == "auto":
                if os.path.isfile(source):
                    source_type = "file"
                elif source.startswith(('http://', 'https://')):
                    source_type = "url"
                else:
                    source_type = "text"
            
            # Process based on source type
            if source_type == "file":
                # Validate file
                file_size = os.path.getsize(source)
                validation = self.validate_file(source, file_size)
                
                if not validation['is_valid']:
                    processing_result['error'] = validation['error_message']
                    return processing_result
                
                # Extract content based on file type
                file_type = validation['file_type']
                file_ext = validation['file_info']['extension']
                
                if file_type == 'text':
                    content = self.extract_text_from_txt(source)
                elif file_ext == '.pdf':
                    content = self.extract_text_from_pdf(source)
                elif file_ext in ['.docx', '.doc']:
                    content = self.extract_text_from_docx(source)
                elif file_ext == '.json':
                    content = self.extract_text_from_json(source)
                elif file_ext == '.csv':
                    content = self.extract_text_from_csv(source)
                else:
                    raise Exception(f"Unsupported file format: {file_ext}")
                
                processing_result['metadata']['file_info'] = validation['file_info']
            
            elif source_type == "url":
                content = self.extract_content_from_url(source)
                processing_result['metadata']['source_url'] = source
            
            elif source_type == "text":
                content = source.strip()
            
            else:
                raise Exception(f"Unknown source type: {source_type}")
            
            # Validate content length
            if len(content.strip()) < 10:
                processing_result['error'] = "Content too short. Minimum 10 characters required."
                return processing_result
            
            # Generate content hash for uniqueness
            content_hash = hashlib.md5(content.encode()).hexdigest()
            
            # Update metadata
            processing_result['success'] = True
            processing_result['content'] = content
            processing_result['metadata']['content_hash'] = content_hash
            processing_result['metadata']['word_count'] = len(content.split())
            processing_result['metadata']['character_count'] = len(content)
            processing_result['metadata']['source_type'] = source_type
            
            print(f"✅ Content processed successfully: {processing_result['metadata']['word_count']} words")
            
            return processing_result
        
        except Exception as e:
            processing_result['error'] = str(e)
            print(f"❌ Content processing failed: {e}")
            return processing_result
    
    def generate_questions_from_content(
        self, 
        content: str, 
        topic: str,
        num_questions: int = 10,
        difficulty: str = 'Medium',
        question_types: List[str] = None
    ) -> Dict[str, Any]:
        """
        Generate quiz questions from content using Google Gemini AI
        
        Args:
            content: The text content to generate questions from
            topic: Topic/subject of the content
            num_questions: Number of questions to generate
            difficulty: Difficulty level (Easy, Medium, Hard)
            question_types: List of question types to generate
        
        Returns:
            Dictionary with generated questions and metadata
        """
        
        if not self.gemini_model:
            return {
                'success': False,
                'error': 'Gemini AI not configured. Please set GEMINI_API_KEY.',
                'questions': []
            }
        
        if question_types is None:
            question_types = ['Multiple Choice', 'True/False', 'Short Answer']
        
        try:
            # Chunk content if too long (Gemini has token limits)
            max_content_length = 8000
            if len(content) > max_content_length:
                logger.warning(f"Content too long ({len(content)} chars). Truncating to {max_content_length} chars.")
                content = content[:max_content_length] + "\n\n[Content truncated for processing]"
            
            # Create detailed prompt for question generation
            prompt = f"""You are an expert quiz creator. Generate {num_questions} high-quality quiz questions based on the following content.

**Topic:** {topic}
**Difficulty Level:** {difficulty}
**Question Types to Include:** {', '.join(question_types)}

**Content:**
{content}

**Requirements:**
1. Generate exactly {num_questions} questions
2. Mix question types: {', '.join(question_types)}
3. Difficulty should be {difficulty} level
4. Each question must be clear and unambiguous
5. For Multiple Choice questions: provide 4 options (A, B, C, D) with only ONE correct answer
6. For True/False questions: provide a statement that is clearly true or false
7. For Short Answer questions: provide a concise correct answer
8. Include detailed explanations for each answer
9. Ensure questions test comprehension, not just memorization

**Output Format (STRICT JSON):**
{{
  "questions": [
    {{
      "question_text": "Question here?",
      "question_type": "Multiple Choice" | "True/False" | "Short Answer",
      "options": ["A) Option 1", "B) Option 2", "C) Option 3", "D) Option 4"],  // Only for Multiple Choice
      "correct_answer": "Correct answer here",
      "explanation": "Detailed explanation of why this is correct",
      "difficulty": "{difficulty}",
      "topic": "{topic}"
    }}
  ]
}}

Generate ONLY valid JSON. Do not include any text before or after the JSON."""

            logger.info(f"🤖 Generating {num_questions} questions using Gemini AI...")
            
            # Call Gemini AI
            response = self.gemini_model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Extract JSON from response (sometimes Gemini adds markdown code blocks)
            json_match = re.search(r'```(?:json)?\s*(\{.*\})\s*```', response_text, re.DOTALL)
            if json_match:
                response_text = json_match.group(1)
            
            # Parse JSON response
            questions_data = json.loads(response_text)
            
            # Validate and process questions
            generated_questions = []
            for q in questions_data.get('questions', []):
                # Ensure all required fields are present
                if 'question_text' in q and 'correct_answer' in q:
                    question = {
                        'question_text': q['question_text'],
                        'question_type': q.get('question_type', 'Multiple Choice'),
                        'correct_answer': q['correct_answer'],
                        'explanation': q.get('explanation', 'No explanation provided.'),
                        'difficulty': difficulty,
                        'topic': topic,
                        'options': q.get('options', []),
                        'generated_by': 'Gemini AI',
                        'source': 'PDF Content'
                    }
                    generated_questions.append(question)
            
            logger.info(f"✅ Successfully generated {len(generated_questions)} questions from content")
            
            return {
                'success': True,
                'questions': generated_questions,
                'total_generated': len(generated_questions),
                'metadata': {
                    'topic': topic,
                    'difficulty': difficulty,
                    'question_types': question_types,
                    'generated_at': datetime.now().isoformat(),
                    'ai_model': 'Gemini-1.5-Flash'
                }
            }
        
        except json.JSONDecodeError as e:
            logger.error(f"❌ Failed to parse Gemini response as JSON: {e}")
            logger.error(f"Response: {response_text[:500]}")
            return {
                'success': False,
                'error': 'Failed to parse AI response. Please try again.',
                'questions': [],
                'raw_response': response_text[:500]
            }
        
        except Exception as e:
            logger.error(f"❌ Question generation failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'questions': []
            }
    
    def process_pdf_and_generate_questions(
        self,
        pdf_path: str,
        topic: str,
        num_questions: int = 10,
        difficulty: str = 'Medium',
        question_types: List[str] = None
    ) -> Dict[str, Any]:
        """
        Complete pipeline: Extract text from PDF and generate questions
        
        Args:
            pdf_path: Path to PDF file
            topic: Topic/subject
            num_questions: Number of questions to generate
            difficulty: Difficulty level
            question_types: Types of questions to generate
        
        Returns:
            Dictionary with questions and metadata
        """
        
        result = {
            'success': False,
            'questions': [],
            'metadata': {},
            'error': None
        }
        
        try:
            # Step 1: Extract text from PDF
            logger.info(f"📄 Extracting text from PDF: {pdf_path}")
            content = self.extract_text_from_pdf(pdf_path, use_advanced=True)
            
            if not content or len(content.strip()) < 50:
                result['error'] = 'Insufficient content extracted from PDF'
                return result
            
            # Get content summary
            summary = self.get_content_summary(content)
            result['metadata']['content_summary'] = summary
            
            # Step 2: Generate questions using AI
            logger.info(f"🤖 Generating {num_questions} questions from extracted content...")
            questions_result = self.generate_questions_from_content(
                content=content,
                topic=topic,
                num_questions=num_questions,
                difficulty=difficulty,
                question_types=question_types
            )
            
            if questions_result['success']:
                result['success'] = True
                result['questions'] = questions_result['questions']
                result['metadata'].update(questions_result['metadata'])
                result['metadata']['pdf_path'] = pdf_path
                
                logger.info(f"✅ Successfully generated {len(result['questions'])} questions from PDF")
            else:
                result['error'] = questions_result.get('error', 'Unknown error')
            
            return result
        
        except Exception as e:
            logger.error(f"❌ PDF processing and question generation failed: {e}")
            result['error'] = str(e)
            return result
    
    def get_content_summary(self, content: str) -> Dict[str, Any]:
        """Generate a summary of processed content"""
        words = content.split()
        sentences = re.split(r'[.!?]+', content)
        
        # Extract key topics (simple keyword extraction)
        word_freq = {}
        for word in words:
            word = re.sub(r'[^\w]', '', word.lower())
            if len(word) > 3:  # Only words longer than 3 characters
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords
        top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'character_count': len(content),
            'top_keywords': [kw[0] for kw in top_keywords],
            'estimated_reading_time': max(1, len(words) // 200),  # ~200 words per minute
            'content_preview': content[:200] + "..." if len(content) > 200 else content
        }

# Usage example and testing
if __name__ == "__main__":
    processor = ContentProcessor()
    
    # Test with sample text
    sample_text = "This is a sample text for testing the content processor functionality."
    result = processor.process_content(sample_text, "text")
    
    if result['success']:
        print("✅ Content processing test successful")
        summary = processor.get_content_summary(result['content'])
        print(f"Summary: {summary}")
    else:
        print(f"❌ Test failed: {result['error']}")